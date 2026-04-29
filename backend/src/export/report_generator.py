"""Automated BRSR / GRI report generation from validated data + AI narratives."""
import logging
import os
from datetime import datetime
from typing import Dict, List, Optional
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy.orm import Session

from src.common.models import NormalizedData

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "templates")
OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "data", "outputs")


class ReportGenerator:
    """Fills official BRSR / GRI DOCX templates with validated data + narratives."""

    def __init__(self, db_session: Session):
        self.db = db_session
        os.makedirs(OUTPUTS_DIR, exist_ok=True)

    # ==================================================================
    # BRSR
    # ==================================================================

    def generate_brsr_report(
        self,
        upload_id: UUID,
        narratives: List[Dict],
        company_info: Dict,
    ) -> str:
        template_path = os.path.join(TEMPLATES_DIR, "brsr_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            self._build_brsr_skeleton(doc)

        self._fill_section_a(doc, company_info)
        self._fill_section_b(doc, narratives)
        self._fill_principle_6(doc, narratives, upload_id)

        filename = f"BRSR_Report_{upload_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(OUTPUTS_DIR, filename)
        doc.save(output_path)
        logger.info(f"BRSR report saved: {output_path}")
        return output_path

    # ------------------------------------------------------------------

    def _build_brsr_skeleton(self, doc: Document) -> None:
        """Build a minimal BRSR skeleton when no template file exists."""
        doc.add_heading("BRSR Core — Business Responsibility and Sustainability Report", level=0)

        doc.add_heading("Section A: General Disclosures", level=1)
        for label in [
            "Company Name", "CIN", "Registered Office", "Corporate Office",
            "E-mail", "Telephone", "Website", "Financial Year", "Reporting Boundary",
        ]:
            placeholder = "{{" + "_".join(label.upper().split()) + "}}"
            doc.add_paragraph(f"{label}: {placeholder}")

        doc.add_heading("Section B: Management and Process Disclosures", level=1)
        table_b = doc.add_table(rows=1, cols=2)
        table_b.style = "Table Grid"
        table_b.cell(0, 0).text = "Management and Process Disclosures"
        table_b.cell(0, 1).text = "Disclosure"

        doc.add_heading("Section C: Principle-wise Performance Disclosure", level=1)
        doc.add_heading("PRINCIPLE 6: Environment", level=2)
        table_p6 = doc.add_table(rows=1, cols=3)
        table_p6.style = "Table Grid"
        table_p6.cell(0, 0).text = "PRINCIPLE 6"
        table_p6.cell(0, 1).text = "Essential Indicator"
        table_p6.cell(0, 2).text = "Value"

        for indicator in [
            "Total electricity consumption",
            "Energy from renewable sources",
            "Scope 1 emissions",
            "Scope 2 emissions",
            "Emissions intensity",
            "Total water withdrawal",
            "Hazardous waste generated",
            "Non-hazardous waste generated",
        ]:
            row = table_p6.add_row()
            row.cells[0].text = ""
            row.cells[1].text = indicator
            row.cells[2].text = ""

    # ------------------------------------------------------------------

    def _fill_section_a(self, doc: Document, company_info: Dict) -> None:
        replacements = {
            "{{COMPANY_NAME}}": company_info.get("name", "N/A"),
            "{{CIN}}": company_info.get("cin", "N/A"),
            "{{REGISTERED_OFFICE}}": company_info.get("registered_office", "N/A"),
            "{{CORPORATE_OFFICE}}": company_info.get("corporate_office", "N/A"),
            "{{EMAIL}}": company_info.get("email", "N/A"),
            "{{E-MAIL}}": company_info.get("email", "N/A"),
            "{{TELEPHONE}}": company_info.get("telephone", "N/A"),
            "{{WEBSITE}}": company_info.get("website", "N/A"),
            "{{FINANCIAL_YEAR}}": company_info.get("financial_year", "2024-25"),
            "{{REPORTING_BOUNDARY}}": company_info.get("boundary", "Standalone"),
        }
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

    def _fill_section_b(self, doc: Document, narratives: List[Dict]) -> None:
        mgmt_narratives = [
            n for n in narratives if n.get("section") == "management_approach"
        ]
        for table in doc.tables:
            header_text = table.cell(0, 0).text
            if "Management" in header_text and "Process" in header_text:
                for narrative in mgmt_narratives:
                    self._insert_narrative(
                        table, narrative["indicator"], narrative["content"]
                    )
                break

    def _fill_principle_6(
        self, doc: Document, narratives: List[Dict], upload_id: UUID
    ) -> None:
        data_map = self._load_indicator_map(upload_id)

        total_electricity = self._sum_indicator(data_map, "Total Electricity")
        total_renewable = self._sum_indicator(data_map, "Renewable Energy")
        scope1 = self._sum_indicator(data_map, "Scope 1 Emissions")
        scope2 = self._sum_indicator(data_map, "Scope 2 Emissions")
        production = self._sum_indicator(data_map, "Production Volume")
        water = self._sum_indicator(data_map, "Total Water Consumption")
        haz_waste = self._sum_indicator(data_map, "Hazardous Waste")
        non_haz_waste = self._sum_indicator(data_map, "Non-Hazardous Waste")

        renewable_pct = (
            (total_renewable / total_electricity * 100)
            if total_electricity > 0
            else 0
        )
        intensity = (
            (scope1 + scope2) / production if production > 0 else 0
        )

        fill_map = {
            "total electricity consumption": f"{total_electricity:,.2f} MWh",
            "energy from renewable sources": f"{total_renewable:,.2f} MWh ({renewable_pct:.1f}%)",
            "scope 1 emissions": f"{scope1:,.2f} tonnes CO\u2082e",
            "scope 2 emissions": f"{scope2:,.2f} tonnes CO\u2082e",
            "emissions intensity": f"{intensity:.4f} tonnes CO\u2082e/tonne product",
            "total water withdrawal": f"{water:,.2f} m\u00b3",
            "hazardous waste generated": f"{haz_waste:,.2f} tonnes",
            "non-hazardous waste generated": f"{non_haz_waste:,.2f} tonnes",
        }

        for table in doc.tables:
            cell0 = table.cell(0, 0).text
            if "PRINCIPLE 6" in cell0 or "Environment" in cell0:
                for label, value in fill_map.items():
                    self._fill_table_cell(table, label, value)
                break

    # ==================================================================
    # GRI
    # ==================================================================

    def generate_gri_report(
        self,
        upload_id: UUID,
        narratives: List[Dict],
        company_info: Dict,
    ) -> str:
        template_path = os.path.join(TEMPLATES_DIR, "gri_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            self._build_gri_skeleton(doc)

        self._fill_gri_2(doc, company_info)
        self._fill_gri_302(doc, narratives, upload_id)
        self._fill_gri_305(doc, narratives, upload_id)
        self._fill_gri_303(doc, narratives, upload_id)
        self._fill_gri_306(doc, narratives, upload_id)

        filename = f"GRI_Report_{upload_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(OUTPUTS_DIR, filename)
        doc.save(output_path)
        logger.info(f"GRI report saved: {output_path}")
        return output_path

    # ------------------------------------------------------------------

    def _build_gri_skeleton(self, doc: Document) -> None:
        doc.add_heading("GRI Standards Sustainability Report", level=0)

        doc.add_heading("GRI 2: General Disclosures 2021", level=1)
        for label in [
            "Organization Name", "CIN", "Headquarters", "Email",
            "Website", "Reporting Period", "Reporting Boundary",
        ]:
            placeholder = "{{" + "_".join(label.upper().split()) + "}}"
            doc.add_paragraph(f"{label}: {placeholder}")

        for disclosure, title, indicators in [
            ("302", "Energy", [
                ("302-1", "Electricity consumption", ""),
                ("302-1", "Fuel consumption", ""),
                ("302-1", "Total energy consumption", ""),
                ("302-3", "Energy intensity", ""),
            ]),
            ("305", "Emissions", [
                ("305-1", "Direct (Scope 1) GHG emissions", ""),
                ("305-2", "Energy indirect (Scope 2) GHG emissions", ""),
                ("305-3", "Other indirect (Scope 3) GHG emissions", ""),
                ("305-4", "GHG emissions intensity", ""),
            ]),
            ("303", "Water and Effluents", [
                ("303-3", "Total water withdrawal", ""),
                ("303-4", "Water discharge", ""),
                ("303-5", "Water consumption", ""),
            ]),
            ("306", "Waste", [
                ("306-3", "Hazardous waste generated", ""),
                ("306-3", "Non-hazardous waste generated", ""),
                ("306-4", "Waste diverted from disposal", ""),
                ("306-5", "Waste directed to disposal", ""),
            ]),
        ]:
            doc.add_heading(f"GRI {disclosure}: {title}", level=1)
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            tbl.cell(0, 0).text = f"{disclosure}-1"
            tbl.cell(0, 1).text = "Indicator"
            tbl.cell(0, 2).text = "Value"
            for disc_id, ind_name, _ in indicators:
                row = tbl.add_row()
                row.cells[0].text = disc_id
                row.cells[1].text = ind_name
                row.cells[2].text = ""

    # ------------------------------------------------------------------

    def _fill_gri_2(self, doc: Document, company_info: Dict) -> None:
        replacements = {
            "{{ORGANIZATION_NAME}}": company_info.get("name", "N/A"),
            "{{CIN}}": company_info.get("cin", "N/A"),
            "{{HEADQUARTERS}}": company_info.get("corporate_office", "N/A"),
            "{{EMAIL}}": company_info.get("email", "N/A"),
            "{{WEBSITE}}": company_info.get("website", "N/A"),
            "{{REPORTING_PERIOD}}": company_info.get("financial_year", "2024-25"),
            "{{REPORTING_BOUNDARY}}": company_info.get("boundary", "Operational control"),
        }
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

    def _fill_gri_302(
        self, doc: Document, narratives: List[Dict], upload_id: UUID
    ) -> None:
        data_map = self._load_indicator_map(upload_id)
        electricity = self._sum_indicator(data_map, "electricity")
        fuel = self._sum_indicator(data_map, "fuel")
        electricity_gj = electricity * 3.6  # MWh → GJ

        for table in doc.tables:
            if self._table_contains(table, "302"):
                self._fill_table_cell(table, "Electricity consumption", f"{electricity_gj:,.2f} GJ")
                self._fill_table_cell(table, "Fuel consumption", f"{fuel:,.2f} GJ")
                self._fill_table_cell(table, "Total energy consumption", f"{electricity_gj + fuel:,.2f} GJ")
                if electricity > 0:
                    production = self._sum_indicator(data_map, "production")
                    if production > 0:
                        self._fill_table_cell(
                            table, "Energy intensity",
                            f"{(electricity_gj + fuel) / production:,.4f} GJ/tonne product",
                        )
                break

        mgmt = next(
            (n for n in narratives
             if "electricity" in n.get("indicator", "").lower()
             and n.get("section") == "management_approach"),
            None,
        )
        if mgmt:
            self._insert_section_narrative(doc, "302-1", mgmt["content"])

    def _fill_gri_305(
        self, doc: Document, narratives: List[Dict], upload_id: UUID
    ) -> None:
        data_map = self._load_indicator_map(upload_id)
        scope1 = self._sum_indicator(data_map, "Scope 1")
        scope2 = self._sum_indicator(data_map, "Scope 2")
        scope3 = self._sum_indicator(data_map, "Scope 3")
        production = self._sum_indicator(data_map, "production")

        for table in doc.tables:
            if self._table_contains(table, "305"):
                self._fill_table_cell(table, "Direct (Scope 1) GHG emissions", f"{scope1:,.2f} tonnes CO\u2082e")
                self._fill_table_cell(table, "Energy indirect (Scope 2) GHG emissions", f"{scope2:,.2f} tonnes CO\u2082e")
                self._fill_table_cell(table, "Other indirect (Scope 3) GHG emissions", f"{scope3:,.2f} tonnes CO\u2082e")
                if production > 0:
                    intensity = (scope1 + scope2) / production
                    self._fill_table_cell(table, "GHG emissions intensity", f"{intensity:,.4f} tonnes CO\u2082e/tonne product")
                break

    def _fill_gri_303(
        self, doc: Document, narratives: List[Dict], upload_id: UUID
    ) -> None:
        data_map = self._load_indicator_map(upload_id)
        water = self._sum_indicator(data_map, "water")
        discharge = self._sum_indicator(data_map, "discharge")
        consumption = water - discharge if water > discharge else water

        for table in doc.tables:
            if self._table_contains(table, "303"):
                self._fill_table_cell(table, "Total water withdrawal", f"{water:,.2f} m\u00b3")
                self._fill_table_cell(table, "Water discharge", f"{discharge:,.2f} m\u00b3")
                self._fill_table_cell(table, "Water consumption", f"{consumption:,.2f} m\u00b3")
                break

    def _fill_gri_306(
        self, doc: Document, narratives: List[Dict], upload_id: UUID
    ) -> None:
        data_map = self._load_indicator_map(upload_id)
        hazardous = self._sum_indicator(data_map, "hazardous")
        non_haz = self._sum_indicator(data_map, "non-hazardous")
        diverted = self._sum_indicator(data_map, "diverted")
        disposed = self._sum_indicator(data_map, "disposal")

        for table in doc.tables:
            if self._table_contains(table, "306"):
                self._fill_table_cell(table, "Hazardous waste generated", f"{hazardous:,.2f} tonnes")
                self._fill_table_cell(table, "Non-hazardous waste generated", f"{non_haz:,.2f} tonnes")
                self._fill_table_cell(table, "Waste diverted from disposal", f"{diverted:,.2f} tonnes")
                self._fill_table_cell(table, "Waste directed to disposal", f"{disposed:,.2f} tonnes")
                break

    # ==================================================================
    # Reconciliation notes (BRSR vs GRI)
    # ==================================================================

    @staticmethod
    def generate_reconciliation_notes(
        brsr_data: Optional[Dict] = None,
        gri_data: Optional[Dict] = None,
    ) -> List[Dict]:
        return [
            {
                "topic": "Scope 2 GHG Emissions",
                "brsr_approach": "Location-based method (grid average emission factor)",
                "gri_approach": "Both location-based and market-based methods",
                "reconciliation": (
                    "BRSR report uses location-based only. GRI report includes both. "
                    "Values may differ with renewable energy certificates or PPAs."
                ),
                "impact": "Potential difference in reported Scope 2 emissions",
            },
            {
                "topic": "Energy Consumption Units",
                "brsr_approach": "Reports in MWh",
                "gri_approach": "Requires GJ (gigajoules)",
                "reconciliation": "Conversion: 1 MWh = 3.6 GJ",
                "impact": "Same energy, different units",
            },
            {
                "topic": "Organizational Boundary",
                "brsr_approach": "Standalone or consolidated per financial reporting",
                "gri_approach": "Operational control, financial control, or equity share",
                "reconciliation": "Both reports use operational control approach for consistency",
                "impact": "No material difference if same boundary applied",
            },
            {
                "topic": "Reporting Period",
                "brsr_approach": "Financial year (April 1 \u2013 March 31 for India)",
                "gri_approach": "Any 12-month period",
                "reconciliation": "Both reports cover the same FY",
                "impact": "No difference \u2014 same period",
            },
        ]

    # ==================================================================
    # Shared helpers
    # ==================================================================

    def _load_indicator_map(self, upload_id: UUID) -> Dict[str, List]:
        """Load NormalizedData grouped by indicator name."""
        rows = (
            self.db.query(NormalizedData)
            .filter(NormalizedData.upload_id == upload_id)
            .all()
        )
        indicator_map: Dict[str, List] = {}
        for r in rows:
            name = r.indicator.matched_indicator if r.indicator else str(r.indicator_id)
            indicator_map.setdefault(name, []).append(r)
        return indicator_map

    @staticmethod
    def _sum_indicator(data_map: Dict[str, List], keyword: str) -> float:
        total = 0.0
        kw = keyword.lower()
        for name, records in data_map.items():
            if kw in name.lower():
                total += sum(r.normalized_value for r in records)
        return total

    @staticmethod
    def _fill_table_cell(table, indicator_label: str, value: str) -> None:
        label_lower = indicator_label.lower()
        for row in table.rows:
            for cell in row.cells:
                if label_lower in cell.text.lower():
                    value_cell = row.cells[-1]
                    value_cell.text = value
                    for p in value_cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(11)
                            run.font.name = "Calibri"
                    return

    @staticmethod
    def _insert_narrative(table, indicator: str, content: str) -> None:
        ind_lower = indicator.lower()
        for row in table.rows:
            if ind_lower in row.cells[0].text.lower():
                cell = row.cells[-1]
                cell.text = content
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = "Calibri"
                return
        row = table.add_row()
        row.cells[0].text = indicator
        row.cells[-1].text = content

    @staticmethod
    def _insert_section_narrative(doc: Document, disclosure_id: str, content: str) -> None:
        for i, paragraph in enumerate(doc.paragraphs):
            if disclosure_id in paragraph.text:
                new_para = doc.add_paragraph(content)
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in new_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"
                return

    @staticmethod
    def _table_contains(table, text: str) -> bool:
        try:
            return text in table.cell(0, 0).text
        except Exception:
            return False
